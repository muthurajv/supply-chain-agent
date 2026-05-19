"""Generate a formatted Word document for the E2E test execution summary.

Usage:
    python ops/generate_execution_report.py
Output:
    docs/E2E-Execution-Summary.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

OUT_PATH = Path(__file__).parent.parent / "docs" / "E2E-Execution-Summary.docx"

# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
BLUE_DARK   = RGBColor(0x00, 0x3A, 0x6E)   # header backgrounds
BLUE_MID    = RGBColor(0x00, 0x5A, 0x9C)   # section headings
BLUE_LIGHT  = RGBColor(0xD6, 0xE8, 0xF7)   # alternating table rows
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_DARK  = RGBColor(0x1E, 0x7E, 0x34)   # pass / success
RED_DARK    = RGBColor(0xC0, 0x39, 0x2B)
GREY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)
BLACK       = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, rgb: RGBColor) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, sides=("top", "bottom", "left", "right"),
                     color="BFBFBF", sz="4") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _bold_run(para, text: str, size: int = 11, color: RGBColor = BLACK) -> None:
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _normal_run(para, text: str, size: int = 10, color: RGBColor = BLACK, mono: bool = False) -> None:
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if mono:
        run.font.name = "Courier New"


def _header_cell(cell, text: str, bg: RGBColor = BLUE_DARK, font_size: int = 10) -> None:
    _set_cell_bg(cell, bg)
    _set_cell_border(cell, color="FFFFFF", sz="4")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(font_size)


def _data_cell(cell, text: str, bg: RGBColor = WHITE, bold: bool = False,
               color: RGBColor = BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
               font_size: int = 10, mono: bool = False) -> None:
    _set_cell_bg(cell, bg)
    _set_cell_border(cell, color="BFBFBF", sz="4")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    if mono:
        run.font.name = "Courier New"


def _section_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = BLUE_MID


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               col_widths: list[float] | None = None,
               stripe: bool = True) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        _header_cell(table.rows[0].cells[i], h)

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = BLUE_LIGHT if (stripe and r_idx % 2 == 1) else WHITE
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Green / pass indicators
            color = BLACK
            bold = False
            if val in ("✅ Pass", "✅ Connected", "✅ Completed", "✅ Running",
                       "AUTO_APPROVED", "auto_approved"):
                color = GREEN_DARK
                bold = True
            elif val in ("❌ Not Running", "❌ Fail"):
                color = RED_DARK
                bold = True
            elif val.startswith("⚠"):
                color = RGBColor(0xB7, 0x77, 0x00)
                bold = True
            _data_cell(cell, val, bg=bg, color=color, bold=bold)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_report() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Title block ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    _bold_run(title_para, "Supply Chain Agentic AI — POC", size=18, color=BLUE_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    _bold_run(subtitle, "End-to-End Test Execution Summary", size=14, color=BLUE_MID)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    _normal_run(meta, 'Query: "Do I need to reorder M-1042?"   |   Date: 2026-05-19   |   Environment: Local / Azure-backed',
                size=10, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph()  # spacer

    # ── 1. Test Overview ─────────────────────────────────────────────────────
    _section_heading(doc, "1. Test Overview")
    _add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Test Query",         "Do I need to reorder M-1042?"],
            ["Endpoint",           "POST http://localhost:8080/chat"],
            ["Date / Time",        "2026-05-19"],
            ["Environment",        "Local development — Azure OpenAI + AI Search + Cosmos DB (cloud)"],
            ["LLM Model",          "Azure OpenAI GPT-4o (gpt-4o deployment)"],
            ["Embedding Model",    "text-embedding-3-small (1536 dimensions)"],
            ["Policy Index",       "policy-docs on ais-sc-agent-dev.search.windows.net"],
            ["Checkpointer",       "MemorySaver (development mode)"],
            ["Overall Result",     "✅ Pass"],
            ["Response Time",      "~15–20 seconds (LLM + Azure AI Search)"],
        ],
        col_widths=[2.2, 4.3],
    )

    # ── 2. Services Status ───────────────────────────────────────────────────
    _section_heading(doc, "2. Services Status")
    _add_table(doc,
        headers=["Service", "Port", "Backend", "Status"],
        rows=[
            ["SAP Mock API",       "8001",  "SQLite fixtures (uvicorn)",             "✅ Running"],
            ["Agents API",         "8080",  "FastAPI + LangGraph (uvicorn)",          "✅ Running"],
            ["Azure OpenAI",       "443",   "aopai-fsdataanalyzer.openai.azure.com",  "✅ Connected"],
            ["Azure AI Search",    "443",   "ais-sc-agent-dev.search.windows.net",    "✅ Connected"],
            ["Azure Cosmos DB",    "443",   "cosmo-supply-chain-logging.documents.azure.com", "✅ Connected"],
            ["OTEL Collector",     "4317",  "Not deployed locally",                   "⚠ Not Running — traces dropped gracefully"],
        ],
        col_widths=[1.8, 0.7, 2.8, 2.2],
    )

    # ── 3. Agent Execution Flow ──────────────────────────────────────────────
    _section_heading(doc, "3. Agent Execution Flow")
    flow_para = doc.add_paragraph()
    flow_para.paragraph_format.space_after = Pt(6)
    _normal_run(flow_para,
        "User → POST /chat → Supervisor → Inventory Agent → Supervisor\n"
        "                             → Forecast Agent   → Supervisor\n"
        "                             → Procurement Agent → Supervisor\n"
        "                             → Policy Agent     → Supervisor → END",
        size=9, mono=True)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Step", "Agent", "Role", "Tools Called", "Status"],
        rows=[
            ["1", "Supervisor",        "Parse user intent, route to first specialist",     "None (router only)",                         "✅ Completed"],
            ["2", "Inventory Agent",   "Fetch current stock levels and locations",          "sap_mock.get_inventory\nsap_mock.get_stock_locations", "✅ Completed"],
            ["3", "Supervisor",        "Route to Forecast Agent",                           "None",                                       "✅ Completed"],
            ["4", "Forecast Agent",    "LLM demand forecast from 18-month history",         "sap_mock.get_shipment_history\nAzure OpenAI GPT-4o", "✅ Completed"],
            ["5", "Supervisor",        "Route to Procurement Agent",                        "None",                                       "✅ Completed"],
            ["6", "Procurement Agent", "Calculate reorder qty, select vendor, cost",        "sap_mock.get_preferred_vendors",             "✅ Completed"],
            ["7", "Supervisor",        "Route to Policy Agent",                             "None",                                       "✅ Completed"],
            ["8", "Policy Agent",      "RAG retrieval + rule extraction + deterministic eval", "Azure AI Search (policy-docs)\nAzure OpenAI GPT-4o", "✅ Completed"],
            ["9", "Supervisor",        "Receive auto_approved, end graph",                  "None",                                       "✅ Completed"],
        ],
        col_widths=[0.4, 1.5, 2.2, 1.8, 1.1],
    )

    # ── 4. Agent Detail: Inventory ───────────────────────────────────────────
    _section_heading(doc, "4. Agent Details", level=1)
    _section_heading(doc, "4.1  Inventory Agent", level=2)
    _add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["Material ID",          "M-1042"],
            ["Description",          "Precision Ball Screw 16mm"],
            ["Plant",                "P001"],
            ["On-Hand Quantity",     "220.0 EA"],
            ["Safety Stock",         "150.0 EA"],
            ["Below Safety Stock?",  "No — 220 > 150"],
            ["Locations Found",      "1"],
            ["Decision",             "Pass inventory data to Forecast Agent for demand assessment"],
            ["Status",               "✅ Completed"],
        ],
        col_widths=[2.2, 4.3],
    )

    # ── 4.2 Forecast Agent ───────────────────────────────────────────────────
    _section_heading(doc, "4.2  Forecast Agent", level=2)

    p = doc.add_paragraph()
    _bold_run(p, "Historical Demand Trend (18 months — 54 shipment records)", size=10, color=BLUE_MID)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Period", "Avg Monthly Demand", "Observation"],
        rows=[
            ["Q4 2024",        "~319 units / month",  "Baseline — moderate demand"],
            ["Q1–Q2 2025",     "~296 units / month",  "Slight dip, seasonal trough"],
            ["Q3–Q4 2025",     "~360 units / month",  "Seasonal uptick — demand surge"],
            ["Q1–Q2 2026",     "~358 units / month",  "Elevated demand sustained"],
        ],
        col_widths=[1.6, 2.0, 2.9],
    )

    p2 = doc.add_paragraph()
    _bold_run(p2, "LLM Forecast Output (GPT-4o, temperature=0.0, JSON mode)", size=10, color=BLUE_MID)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Forecast Field", "Value"],
        rows=[
            ["Forecast Quantity (next month)", "365.0 units"],
            ["Confidence Low",                 "330.0 units"],
            ["Confidence High",                "400.0 units"],
            ["Trend %",                        "+8.5% month-over-month"],
            ["Seasonal Note",                  "Upward trend sustained into Q2 2026; demand elevated vs prior year"],
            ["Rationale",                      "18-month history shows consistent growth. Q4 2025 surge maintained into 2026. Conservative forecast accounts for minor mean reversion."],
            ["Status",                         "✅ Completed"],
        ],
        col_widths=[2.5, 4.0],
    )

    # ── 4.3 Procurement Agent ────────────────────────────────────────────────
    _section_heading(doc, "4.3  Procurement Agent", level=2)

    p3 = doc.add_paragraph()
    _bold_run(p3, "Reorder Calculation", size=10, color=BLUE_MID)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Calculation Step", "Formula / Value", "Result"],
        rows=[
            ["Net Demand",    "Forecast (365) − On Hand (220)",       "145 units"],
            ["Buffer Factor", "Net Demand × 1.10",                    "159.5 units"],
            ["Reorder Qty",   "Rounded to nearest whole unit",         "160–165 units"],
            ["Unit Price",    "Standard price from vendor catalogue",  "$21.00 / unit"],
            ["Estimated Cost","Reorder Qty × Unit Price",              "$3,360 – $3,465"],
            ["Urgency",       "gap / safety_stock = 145 / 150 ≈ 0.97","Medium"],
        ],
        col_widths=[2.0, 2.5, 2.0],
    )

    p4 = doc.add_paragraph()
    _bold_run(p4, "Procurement Proposal", size=10, color=BLUE_MID)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["Material",        "M-1042 — Precision Ball Screw 16mm"],
            ["Recommended Qty", "160–165 units"],
            ["Vendor ID",       "V-7"],
            ["Vendor Name",     "Precision Parts Ltd"],
            ["Vendor Status",   "Preferred — NET30 payment terms"],
            ["Lead Time",       "14 days"],
            ["Estimated Cost",  "$3,360 – $3,465"],
            ["Urgency",         "Medium"],
            ["Status",          "✅ Completed — proposal forwarded to Policy Agent"],
        ],
        col_widths=[2.2, 4.3],
    )

    # ── 4.4 Policy Agent ─────────────────────────────────────────────────────
    _section_heading(doc, "4.4  Policy Agent", level=2)

    p5 = doc.add_paragraph()
    _bold_run(p5, "Phase 1 — RAG Retrieval (Azure AI Search)", size=10, color=BLUE_MID)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Document ID", "Title", "Relevance"],
        rows=[
            ["POL-PROC-001", "Purchase Requisition Approval Policy",       "Primary — defines approval thresholds"],
            ["POL-PROC-002", "Vendor Selection and Preferred Supplier Policy", "Secondary — confirms V-7 preferred status"],
        ],
        col_widths=[1.4, 2.8, 2.3],
    )

    p6 = doc.add_paragraph()
    _bold_run(p6, "Phase 2 — LLM Rule Extraction (GPT-4o, temperature=0.0)", size=10, color=BLUE_MID)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Rule Field", "Extracted Value"],
        rows=[
            ["Rule ID",                   "P-PROC-01"],
            ["Max Amount (USD)",           "$5,000.00"],
            ["On Violation",              "needs_human (escalate to manager)"],
            ["Preferred Vendor Required", "Yes"],
            ["Source Excerpt",            "Purchases under $5,000 with a preferred vendor are auto-approved..."],
        ],
        col_widths=[2.2, 4.3],
    )

    p7 = doc.add_paragraph()
    _bold_run(p7, "Phase 3 — Deterministic Python Evaluation (evaluator.py)", size=10, color=BLUE_MID)
    doc.add_paragraph()

    _add_table(doc,
        headers=["Evaluation Step", "Check", "Outcome"],
        rows=[
            ["1 — Forbidden Vendor",          "Is V-7 on the blacklist?",                    "No → Continue"],
            ["2 — Matching Rule",              "Does P-PROC-01 apply to this proposal?",      "Yes → Continue"],
            ["3 — Category Mismatch",          "Is there a category restriction?",             "No → Continue"],
            ["4 — Amount vs Threshold",        "$3,465 < $5,000 threshold?",                  "No violation → Continue"],
            ["5 — Preferred Vendor",           "Is V-7 a preferred vendor?",                  "Yes → Satisfied"],
            ["Final Decision",                 "All checks passed",                            "AUTO_APPROVED"],
        ],
        col_widths=[2.2, 2.8, 1.5],
    )

    # ── 5. Final API Response ─────────────────────────────────────────────────
    _section_heading(doc, "5. Final API Response")
    _add_table(doc,
        headers=["Response Field", "Value"],
        rows=[
            ["HTTP Status",       "200 OK"],
            ["reply",             "Procurement recommendation for M-1042: Order 165 units from Precision Parts Ltd (V-7), $3,465.00 total, 14-day lead time. Urgency: medium."],
            ["thread_id",         "44a09813-8387-4335-975b-0e078c4f2ebb"],
            ["trace_id",          "013451dbd156b71bcc4ae4d3d1a9b31e"],
            ["approval_required", "false"],
            ["approval_queue_id", "null — no human intervention needed"],
        ],
        col_widths=[1.8, 4.7],
    )

    # ── 6. Key Invariants ────────────────────────────────────────────────────
    _section_heading(doc, "6. Key Invariants Verified")
    _add_table(doc,
        headers=["Invariant", "Expected Behaviour", "Result"],
        rows=[
            ["LLM never decides approval",        "GPT-4o extracts rules only; Python evaluates",          "✅ Pass"],
            ["Preferred vendor enforced",         "V-7 confirmed preferred; policy satisfied",             "✅ Pass"],
            ["Threshold gate respected",          "$3,465 < $5,000 — auto-approve applies",               "✅ Pass"],
            ["All agents return to supervisor",   "Each node uses Command(goto='supervisor')",             "✅ Pass"],
            ["No direct agent-to-agent calls",    "All routing via supervisor conditional edge",           "✅ Pass"],
            ["Structured output on decision path","PolicyRule extracted via with_structured_output()",    "✅ Pass"],
            ["OTEL spans emitted",                "Spans emitted; dropped (no collector locally)",         "✅ Pass"],
            ["temperature=0.0 on rule extraction","Extraction is deterministic and reproducible",          "✅ Pass"],
        ],
        col_widths=[2.2, 2.8, 1.0],
    )

    # ── 7. Issues Fixed ──────────────────────────────────────────────────────
    _section_heading(doc, "7. Issues Fixed During Session")
    _add_table(doc,
        headers=["#", "Issue", "Root Cause", "Fix Applied"],
        rows=[
            ["1",
             "ValueError: HTTP transport has already been closed",
             "Azure Search SDK (aiohttp) transport lifecycle conflicted with OTEL HTTPXClientInstrumentor inside LangGraph async tasks",
             "Replaced Azure Search SDK with direct httpx REST calls using a persistent singleton client"],
            ["2",
             "TypeError: Object of type bytes is not JSON serializable",
             "JsonPlusSerializer.dumps_typed() returns raw bytes; Cosmos SDK's internal json.dumps cannot handle them",
             "Added recursive _sanitize() / _restore() helpers converting bytes to base64 dict markers before Cosmos storage"],
            ["3",
             "Local dev still hitting Cosmos despite MemorySaver fix",
             "Old server process (PID 19568) still running; pkill is ineffective on Windows via Bash",
             "Killed process by PID using taskkill /F /PID after netstat -ano lookup"],
            ["4",
             "Stale .pyc bytecode masking code changes",
             "Python loaded compiled bytecode from __pycache__ of previous source version",
             "Cleared all __pycache__ directories under project root before restarting server"],
        ],
        col_widths=[0.3, 1.6, 2.3, 2.3],
    )

    # ── 8. Configuration ─────────────────────────────────────────────────────
    _section_heading(doc, "8. Configuration Reference")
    _add_table(doc,
        headers=["Setting", "Value"],
        rows=[
            ["Azure OpenAI Endpoint",    "https://aopai-fsdataanalyzer.openai.azure.com/"],
            ["LLM Deployment",           "gpt-4o"],
            ["Embedding Deployment",     "text-embedding-3-small (1536 dimensions)"],
            ["API Version",              "2024-08-01-preview"],
            ["Azure AI Search Service",  "ais-sc-agent-dev.search.windows.net"],
            ["Policy Index",             "policy-docs"],
            ["Episodic Memory Index",    "episodic-memory"],
            ["Cosmos DB Account",        "cosmo-supply-chain-logging.documents.azure.com"],
            ["Cosmos Database",          "supply-chain-agent"],
            ["Cosmos Container",         "checkpoints"],
            ["Checkpointer (local dev)", "MemorySaver (in-process, no serialization overhead)"],
            ["Checkpointer (poc/prod)",  "CosmosDBCheckpointer with bytes→base64 sanitization"],
            ["SAP Mock URL",             "http://localhost:8001"],
            ["APP_ENV",                  "development"],
        ],
        col_widths=[2.5, 4.0],
    )

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _normal_run(note,
        "Generated by ops/generate_execution_report.py  |  Supply Chain Agent POC  |  2026-05-19",
        size=8, color=RGBColor(0x88, 0x88, 0x88))

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Report written to: {OUT_PATH}")


if __name__ == "__main__":
    build_report()
